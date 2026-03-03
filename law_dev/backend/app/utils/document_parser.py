import os
import logging
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """Extract text from various document formats"""
    
    try:
        ext = file_extension.lower()
        
        if ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return extract_text_from_docx(file_path)
        elif ext == '.txt':
            return extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
            
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        raise


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using pdfplumber (better) with PyPDF2 fallback"""
    
    # Try pdfplumber first (better for complex PDFs)
    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        if text_parts:
            logger.info(f"Extracted text from PDF using pdfplumber ({len(text_parts)} pages)")
            return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
    
    # Fallback to PyPDF2
    try:
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        if text_parts:
            logger.info(f"Extracted text from PDF using PyPDF2 ({len(text_parts)} pages)")
            return "\n\n".join(text_parts)
        else:
            raise ValueError("No text could be extracted from PDF")
    except Exception as e:
        logger.error(f"PyPDF2 also failed: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    
    try:
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_parts.append(row_text)
        
        if text_parts:
            logger.info("Extracted text from DOCX file")
            return "\n\n".join(text_parts)
        else:
            raise ValueError("No text could be extracted from DOCX")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        if text.strip():
            logger.info("Extracted text from TXT file")
            return text
        else:
            raise ValueError("TXT file is empty")
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            logger.info("Extracted text from TXT file (latin-1 encoding)")
            return text
        except Exception as e:
            logger.error(f"Error reading TXT file with different encoding: {str(e)}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        raise Exception(f"Failed to extract text from TXT: {str(e)}")

